from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify, session
import json
from flask_cors import CORS
from dotenv import load_dotenv
import psycopg2
import secrets
from psycopg2 import OperationalError
from docx import Document
from io import BytesIO
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random
from flask_mail import Mail, Message
import os
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash
import jwt
from datetime import datetime
from psycopg2.extras import RealDictCursor
import re

load_dotenv()

app = Flask(__name__)
app.secret_key = secrets.token_hex(16)
# print(app.secret_key)

UPLOAD_FOLDER = 'static/uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Email Configuration
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER')
app.config['MAIL_PORT'] = os.getenv('MAIL_PORT')
app.config['MAIL_USE_SSL'] = os.getenv('MAIL_USE_SSL') == 'True'
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER')
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
mail = Mail(app)

# Database configuration
DB_CONFIG = {
    "host": os.getenv("DB_HOST", "localhost"),
    "dbname": os.getenv("DB_DATABASE"),
    "user": os.getenv("DB_USER"),
    "password": os.getenv("DB_PASSWORD"),
    "port": os.getenv("DB_PORT", 5432)
}

def check_access():
    return 'user_id' in session

# Function to establish a PostgreSQL database connection
def create_connection():
    try:
        return psycopg2.connect(**DB_CONFIG)
    except OperationalError as e:
        print(f"PostgreSQL connection error: {e}")
        return None


def create_database():
    try:
        # Connect to the default 'postgres' database
        connection = psycopg2.connect(
            dbname="postgres",
            user=os.getenv('DB_USER'),
            password=os.getenv('DB_PASSWORD'),
            host=os.getenv('DB_HOST', 'localhost'),
            cursor_factory=RealDictCursor
        )
        connection.autocommit = True
        cursor = connection.cursor()

        # Check if database already exists
        cursor.execute(f"SELECT 1 FROM pg_database WHERE datname = %s", (os.getenv('DB_DATABASE'),))
        exists = cursor.fetchone()

        if not exists:
            cursor.execute(f"CREATE DATABASE {os.getenv('DB_DATABASE')}")
            print(f"Database '{os.getenv('DB_DATABASE')}' created.")
        else:
            print(f"Database '{os.getenv('DB_DATABASE')}' already exists.")

        cursor.close()
        connection.close()
    except Exception as e:
        print(f"Error occurred while creating PostgreSQL database: {e}")


# Function to initialize tables
def initialize_database():
    connection = create_connection()
    if connection:
        try:
            cursor = connection.cursor()

            cursor.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY,
                email VARCHAR(255) UNIQUE NOT NULL,
                username VARCHAR(255) UNIQUE NOT NULL,
                password VARCHAR(255) NOT NULL,
                role VARCHAR(50) NOT NULL,
                otp VARCHAR(50),
                status VARCHAR(20) DEFAULT 'pending' CHECK (status IN ('pending', 'approved', 'rejected')),
                is_verified BOOLEAN DEFAULT FALSE
            );
            """)

            cursor.execute("""
            CREATE TABLE IF NOT EXISTS subjects (
                subject_id SERIAL PRIMARY KEY,
                subject_name VARCHAR(100) NOT NULL,
                branch VARCHAR(100) NOT NULL,
                semester INT NOT NULL,
                UNIQUE (subject_name, branch)
            );
            """)

            cursor.execute("""
            CREATE TABLE IF NOT EXISTS questions (
                question_id SERIAL PRIMARY KEY,
                question_text TEXT NOT NULL,
                rbt_level INT NOT NULL,
                co INT NOT NULL,
                pi INT NOT NULL,
                marks INT NOT NULL,
                subject_id INT NOT NULL,
                user_id INT NOT NULL,
                image_path VARCHAR(255),
                FOREIGN KEY (subject_id) REFERENCES subjects(subject_id) ON DELETE CASCADE,
                FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
            );
            """)

            connection.commit()
            print("PostgreSQL tables created successfully!")
        except Exception as e:
            print(f"Error creating tables: {e}")
        finally:
            cursor.close()
            connection.close()


 # Routes
@app.route('/')
def index():
    return render_template('index.html')


# ✅ Function to generate a JWT token
def generate_verification_token(email):
    return jwt.encode(
        {'email': email, 'exp': datetime.datetime.utcnow() + datetime.timedelta(hours=1)}, 
        app.secret_key, algorithm="HS256"
    )

# ✅ Function to send a verification email
def send_verification_email(email):
    token = generate_verification_token(email)
    verification_url = url_for('verify_email', token=token, _external=True)

    # 👇 Pass the generated URL to the email body constructor
    text_body, html_body = construct_verification_email(verification_url)

    msg = Message(
        subject="Confirm your PaperGen account",
        recipients=[email]
    )
    msg.body = text_body
    msg.html = html_body

    mail.send(msg)


def is_valid_password(password):
    """Check if the password meets the security criteria"""
    pattern = re.compile(r"^(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{5,}$")
    return pattern.match(password)

def fetch_random_admin_email():
    connection = None
    cursor = None

    try:
        # Create database connection
        connection = create_connection()
        cursor = connection.cursor(dictionary=True)

        # Query to fetch all approved admins
        cursor.execute("SELECT email FROM users WHERE role = 'Admin' AND status = 'approved'")
        admins = cursor.fetchall()

        if admins:
            # Randomly select an admin email
            random_admin = random.choice(admins)
            return random_admin['email']
        else:
            # Return None if no admins are available
            return None

    except Exception as e:
        print(f"Error fetching random admin: {e}")
        return None

    finally:
        if cursor:
            cursor.close()
        if connection:
            connection.close()

@app.route('/signup', methods=['GET', 'POST', 'OPTIONS'])
def signup():
    connection = None
    cursor = None

    if request.method == 'GET':
        return render_template('signup.html')  # Regular user signup page

    if request.method == 'POST':
        try:
            data = request.get_json()
            username = data.get('username')
            email = data.get('email')
            password = data.get('password')
            role = data.get('role')  # Here, we expect 'role' to always be 'User' or some non-admin role

            hashed_password = generate_password_hash(password)

            connection = create_connection()
            cursor = connection.cursor(dictionary=True)

            # ✅ Check for existing email
            cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
            if cursor.fetchone():
                return jsonify({'status': 'error', 'message': 'Email already exists. Please use a different email.'})

            # ✅ Check for existing username
            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            if cursor.fetchone():
                return jsonify({'status': 'error', 'message': 'Username is already taken. Please choose a different username.'})

            # ✅ Insert regular user (non-admin)
            cursor.execute("""
                INSERT INTO users (username, email, password, role, is_verified, status)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (username, email, hashed_password, role, False, 'pending'))
            connection.commit()

            # ✅ Notify the admin about the new user signup
            notify_random_admin(username, role, request)

            return jsonify({
                'status': 'success',
                'message': 'Signup successful. Please wait for admin approval.'
            })

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'An error occurred during signup. Please try again. {str(e)}'
            })

        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()


@app.route('/adminsignup', methods=['GET', 'POST', 'OPTIONS'])
def adminsignup():
    connection = None
    cursor = None

    if request.method == 'GET':
        return render_template('adminsignup.html')  # Admin signup page (HTML for admins)

    if request.method == 'POST':
        try:
            data = request.get_json()
            username = data.get('username')
            email = data.get('email')
            password = data.get('password')

            hashed_password = generate_password_hash(password)

            connection = create_connection()
            cursor = connection.cursor(dictionary=True)

            # ✅ Check for existing email
            cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
            if cursor.fetchone():
                return jsonify({'status': 'error', 'message': 'Email already exists. Please use a different email.'})

            # ✅ Check for existing username
            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            if cursor.fetchone():
                return jsonify({'status': 'error', 'message': 'Username is already taken. Please choose a different username.'})

            # ✅ Check if this is the first admin
            cursor.execute("SELECT * FROM users WHERE role = 'Admin' AND status = 'approved' LIMIT 1")
            existing_admin = cursor.fetchone()

            if not existing_admin:  # First admin signup
                cursor.execute("""
                    INSERT INTO users (username, email, password, role, is_verified, status)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (username, email, hashed_password, 'Admin', True, 'approved'))
                connection.commit()

                send_verification_email(email)  # Send verification email to the first admin

                return jsonify({
                    'status': 'success',
                    'message': 'You are the first admin. Your account is approved and a verification email has been sent.'
                })

            else:  # Subsequent admin signup (pending)
                cursor.execute("""
                    INSERT INTO users (username, email, password, role, is_verified, status)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (username, email, hashed_password, 'Admin', False, 'pending'))
                connection.commit()

                send_verification_email(email)  # Send verification email to the new admin
                notify_random_admin(username, 'Admin', request)  # Notify existing admins about the new admin signup

                return jsonify({
                    'status': 'success',
                    'message': 'Signup successful. Your account is pending approval. Please check your inbox.'
                })

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'An error occurred during signup. Please try again. {str(e)}'
            })

        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()


def notify_random_admin(new_username, new_user_role, request):
    connection = create_connection()
    cursor = connection.cursor(dictionary=True)

    try:
        # Get all approved admins
        cursor.execute("SELECT email FROM users WHERE role = 'Admin' AND status = 'approved'")
        admins = cursor.fetchall()

        if admins:
            # Pick one randomly
            selected_admin = random.choice(admins)
            admin_email = selected_admin['email']

            # Generate email content
            subject = f"New {new_user_role} Signup Request"
            text_body, html_body = construct_admin_notification_email(new_username, admin_email, new_user_role)

            # Create message object
            msg = Message(
                subject=subject,
                recipients=[admin_email]
            )
            msg.body = text_body
            msg.html = html_body

            # Send the email
            mail.send(msg)

    finally:
        cursor.close()
        connection.close()

@app.route('/admin_user')
def admin_user():
    # Check if the user is authenticated by verifying if 'user_id' is in the session
    if not check_access():  
        return redirect(url_for('login'))

    # If the user is authenticated, render the admin user page
    return render_template('admin_user.html')



@app.route('/get-users')
def get_users():
    connection = create_connection()
    cursor = connection.cursor()
    cursor.execute("SELECT id, username, email, role, status FROM users")
    users = cursor.fetchall()
    cursor.close()
    connection.close()
    return jsonify(users)


@app.route('/update_user_status', methods=['POST'])
def update_user_status():
    data = request.get_json()
    user_id = data['user_id']
    action = data['action']

    if action == "approve":
        new_status = "approved"
    elif action == "reject":
        new_status = "rejected"
    elif action == "restore":
        new_status = "pending"
    else:
        return jsonify({"error": "Invalid action"}), 400

    connection = create_connection()
    cursor = connection.cursor()

    try:
        # ✅ Update user status
        cursor.execute("UPDATE users SET status = %s WHERE id = %s", (new_status, user_id))
        connection.commit()

        # ✅ Send approval email only if status is approved
        if new_status == "approved":
            cursor.execute("SELECT email, username FROM users WHERE id = %s", (user_id,))
            user = cursor.fetchone()

            if user:
                email, username = user
                send_user_approval_email(email, username)

        return jsonify({"message": f"User status updated to {new_status}."})

    except Exception as e:
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500

    finally:
        cursor.close()
        connection.close()



@app.route('/create_subject', methods=['POST'])
def create_subject():
    # Get the data from the request
    subject_name = request.json.get('subject_name')
    branch = request.json.get('branch')
    semester = request.json.get('semester')
    
    # print(subject_name)
    # print(branch)
    # print(semester)

    try:
        # Establish database connection
        connection = create_connection()
        cursor = connection.cursor()
        
        
        # Check if the subject already exists by subject_name
        cursor.execute("""
            SELECT COUNT(*) FROM subjects WHERE subject_name = %s
        """, (subject_name,))
        result = cursor.fetchone()

        # If the subject already exists, return an error response
        if result[0] > 0:
            return jsonify({"success": False, "message": "Subject already exists."}), 400

        # SQL query to insert the new subject
        cursor.execute("""
            INSERT INTO subjects (subject_name, branch, semester)
            VALUES (%s, %s, %s)
        """, (subject_name, branch, semester))

        # Commit the changes to the database
        connection.commit()

        # Close the cursor and connection
        cursor.close()
        connection.close()

        return jsonify({"success": True, "message": "Subject created successfully!"}), 201

    except Exception as e:
        # Log the error and return a message
        # print(f"Error occurred: {str(e)}")
        return jsonify({"success": False, "message": "Failed to create subject. Please try again later."}), 500




@app.route('/verify-email/<token>')
def verify_email(token):
    try:
        data = jwt.decode(token, app.secret_key, algorithms=["HS256"])
        email = data['email']

        connection = create_connection()
        cursor = connection.cursor()
        cursor.execute("UPDATE users SET is_verified = TRUE WHERE email = %s", (email,))
        connection.commit()
        cursor.close()
        connection.close()

        flash("Email verified successfully! You can now log in.", "success")
    except jwt.ExpiredSignatureError:
        flash("Verification link has expired. Please sign up again.", "danger")
    except jwt.InvalidTokenError:
        flash("Invalid verification link.", "danger")

    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST', 'OPTIONS'])
def login():
    if request.method == 'GET':
        return render_template('login.html')  # ✅ Handle GET request and serve login page

    if request.method == 'POST':
        try:
            data = request.get_json()
            username = data.get('username')
            password = data.get('password')

            if not username or not password:
                return jsonify({
                    'status': 'error',
                    'message': 'Username and password are required.'
                })

            connection = create_connection()
            cursor = connection.cursor()

            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            user = cursor.fetchone()

            if user and check_password_hash(user['password'], password):
                if user['status'] == 'approved':
                    session['user_id'] = user['id']
                    session['username'] = user['username']  # Store the username
                    redirect_url = url_for('add_question') if user['role'] == 'Teacher' else \
                                   url_for('generate_question_paper') if user['role'] == 'Higher Authority' else \
                                   url_for('admin_user')

                    return jsonify({
                        'status': 'success',
                        'message': 'Login successful!',
                        'redirect_url': redirect_url
                    })
                elif user['status'] == 'pending':
                    return jsonify({'status': 'error', 'message': 'Your account is pending approval.'})
                elif user['status'] == 'rejected':
                    return jsonify({'status': 'error', 'message': 'Your account has been rejected.'})
            else:
                return jsonify({'status': 'error', 'message': 'Invalid username or password'})

        except Exception as e:
            return jsonify({'status': 'error', 'message': f'An error occurred: {str(e)}'})

        finally:
            if 'cursor' in locals():
                cursor.close()
            if 'connection' in locals():
                connection.close()


# ✅ Updated login_required decorator
def login_required(f):
    """Decorator to restrict access to logged-in users."""
    def wrapper(*args, **kwargs):
        if 'user_id' not in session:  # ✅ Check user_id instead of username
            return redirect(url_for('login'))  # Redirect to login page
        return f(*args, **kwargs)  # Otherwise, proceed to the page
    wrapper.__name__ = f.__name__
    return wrapper


# ✅ Example of applying the login_required decorator:
@app.route('/protected')
@login_required
def protected_route():
    return "You are logged in!"

def send_user_approval_email(email, username):
    login_url = url_for('login', _external=True)

    # Plain text version
    text_body = f"""
    Hi {username},

    Your account has been approved! 🎉
    You can now log in and explore the platform.

    Login here: {login_url}

    Best regards,  
    The PaperGen Team
    """

    # HTML version
    html_body = f"""
    <html>
    <body>
        <p>Hi <strong>{username}</strong>,</p>
        <p>Your account has been <strong>approved</strong>! 🎉</p>
        <p>You can now <a href="{login_url}" target="_blank">log in</a> and explore the platform.</p>
        <br>
        <p>Best regards,<br><strong>The PaperGen Team</strong></p>
    </body>
    </html>
    """

    # Construct and send the message
    msg = Message(
        subject="Your PaperGen account is now active!",
        recipients=[email],
        body=text_body,
        html=html_body
    )

    mail.send(msg)
    
    
def construct_admin_notification_email(username, email, role):
    # Get the full URL to the admin_user page
    admin_user_url = url_for('admin_user', _external=True)

    # Email subject could be based on the role
    role_display = role.capitalize()

    # Plain text email body
    text_body = f"""
A new {role_display} account has been requested and is currently in a pending state.

Username: {username}
Email: {email}

Please log in to the admin dashboard to approve or reject the request:
{admin_user_url}

Best regards,
Your System
"""

    # HTML email body
    html_body = f"""
<html>
<body>
    <p>A new <strong>{role_display}</strong> account request is currently in a <strong>pending</strong> state.</p>
    <p><strong>Username:</strong> {username}</p>
    <p><strong>Email:</strong> {email}</p>
    <p>Please <a href="{admin_user_url}" target="_blank">log in</a> to the admin dashboard to approve or reject the request.</p>
    <br>
    <p>Best regards,<br>Your System</p>
</body>
</html>
"""

    return text_body, html_body
def construct_verification_email(verification_url):
    current_year = datetime.now().year

    # Plain text version
    text = f"""
Hi,

Welcome to PaperGen! Please verify your email address by clicking the link below:

{verification_url}

⚠️ This link will expire in 1 hour.

If you didn’t create a PaperGen account, you can safely ignore this email.

Thanks,
The PaperGen Team

© {current_year} PaperGen. All rights reserved.
"""

    # HTML version
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-color: #f9f9f9;
                margin: 0;
                padding: 20px;
                color: #333;
            }}
            .container {{
                max-width: 600px;
                background-color: #ffffff;
                padding: 20px;
                border-radius: 8px;
                box-shadow: 0 4px 8px rgba(0,0,0,0.1);
                margin: 0 auto;
            }}
            h1 {{
                color: #4CAF50;
                text-align: center;
            }}
            p {{
                font-size: 16px;
            }}
            .button {{
                display: block;
                width: fit-content;
                background-color: #4CAF50;
                color: #fff;
                padding: 12px 24px;
                border-radius: 5px;
                font-size: 16px;
                font-weight: bold;
                text-align: center;
                text-decoration: none;
                margin: 20px auto;
            }}
            .footer {{
                font-size: 14px;
                color: #777;
                text-align: center;
                margin-top: 30px;
            }}
            .note {{
                font-size: 14px;
                color: #999;
                text-align: center;
                margin-top: 15px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>Verify Your Email</h1>
            <p>Hi there,</p>
            <p>Thanks for signing up for PaperGen! To get started, please confirm your email address by clicking the button below</p>
            <a href="{verification_url}" class="button">Verify Email</a>
            <p class="note">⚠️ This link will expire in 1 hour.</p>
            <p>If you didn’t sign up, you can safely ignore this email.</p>
            <div class="footer">
                © {current_year} PaperGen. All rights reserved.
            </div>
        </div>
    </body>
    </html>
    """

    return text, html

# Function to construct the email content
def construct_email(otp):
    current_year = datetime.now().year
    
    # Plain text version (fallback)
    text = f"""
Hi,

You recently requested to reset your password for your PaperGen account. 
Please use the following OTP to complete the process:

OTP: {otp}

How to reset your password:
1. Enter the OTP on the password reset page.
2. Create a new password for your account.
3. Log in to your account using the new password.

If you didn't request this password reset, please ignore this email or contact our support team immediately.

Thank you for using PaperGen!

© {current_year} PaperGen. All rights reserved.
"""

    # HTML version (for modern email clients)
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-color: #f9f9f9;
                margin: 0;
                padding: 20px;
                color: #333;
                line-height: 1.6;
            }}
            .container {{
                max-width: 600px;
                background-color: #ffffff;
                padding: 20px;
                border-radius: 8px;
                box-shadow: 0 4px 8px rgba(0,0,0,0.1);
                margin: 0 auto;
            }}
            h1 {{
                color: #4CAF50;
                text-align: center;
            }}
            p {{
                font-size: 16px;
                margin-bottom: 20px;
            }}
            .otp {{
                display: block;
                width: fit-content;
                background-color: #4CAF50;
                color: #ffffff;
                padding: 12px 24px;
                border-radius: 5px;
                font-size: 20px;
                font-weight: bold;
                text-align: center;
                margin: 20px auto;
                letter-spacing: 2px;
            }}
            .footer {{
                font-size: 14px;
                color: #777;
                text-align: center;
                margin-top: 30px;
            }}
            a {{
                color: #4CAF50;
                text-decoration: none;
                font-weight: bold;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>PaperGen - Password Reset Request</h1>
            <p>Hello,</p>
            <p>You recently requested to reset your password for your PaperGen account. Please use the OTP below to complete the process:</p>
            <div class="otp">{otp}</div>
            <p><strong>How to reset your password:</strong></p>
            <ol>
                <li>Enter the OTP on the password reset page.</li>
                <li>Create a new password for your account.</li>
                <li>Log in to your account using the new password.</li>
            </ol>
            <p>If you didn't request this password reset, please ignore this email or contact our support team immediately.</p>
            <p>Thank you for using PaperGen!</p>
            <div class="footer">
                © {current_year} PaperGen. All rights reserved. | <a href="https://www.papergen.com">Visit PaperGen</a>
            </div>
        </div>
    </body>
    </html>
    """

    return text, html


@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'GET':
        return render_template('forgot_password.html')
    
    if request.method == 'POST':
        try:
            data = request.get_json()
            email = data.get('email')
            # print(email)

            if not email:
                return jsonify({'error': 'Email is required'}), 400

            connection = create_connection()
            if connection:
                cursor = connection.cursor()
                cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
                user = cursor.fetchone()

                if user:
                    # Generate 6-digit OTP
                    otp = f"{random.randint(100000, 999999)}"

                    # Store OTP in the database
                    cursor.execute("UPDATE users SET otp = %s WHERE email = %s", (otp, email))
                    connection.commit()

                    # Construct email (both text and HTML)
                    text, html = construct_email(otp)

                    # Send the OTP to the user's email
                    msg = Message("PaperGen - Password Reset Request", recipients=[email])
                    msg.body = text
                    msg.html = html
                    mail.send(msg)

                    return jsonify({'message': 'OTP sent to your email address'}), 200
                else:
                    return jsonify({'error': 'Email not found'}), 404
            else:
                return jsonify({'error': 'Failed to connect to the database'}), 500

        except Exception as e:
            return jsonify({'error': f"Database Error: {e}"}), 500
        
        finally:
            if connection:
                cursor.close()
                connection.close()
                
                
@app.route('/verify_otp', methods=['GET', 'POST'])
def verify_otp():
    email = request.args.get('email')

    if request.method == 'GET':
        if not email:
            return "Invalid request. Email is missing.", 400
        
        # Render the OTP verification template
        return render_template('verify_otp.html', email=email)

    if request.method == 'POST':
        try:
            data = request.get_json()
            email = data.get('email')
            otp = data.get('otp')
            new_password = data.get('new_password')

            # Validate OTP and reset password
            connection = create_connection()
            if connection:
                cursor = connection.cursor()
                
                # Get the stored OTP and password
                cursor.execute("SELECT otp, password FROM users WHERE email = %s", (email,))
                stored_data = cursor.fetchone()

                if stored_data and stored_data[0] == otp:
                    stored_password = stored_data[1]

                    # Check if the new password is the same as the previous password
                    if check_password_hash(stored_password, new_password):
                        return jsonify({'error': 'Déjà vu? You can\'t use the same password twice!'}), 400
                    
                    # Hash the new password before storing it
                    hashed_password = generate_password_hash(new_password, method="pbkdf2:sha256", salt_length=8)
                    cursor.execute("UPDATE users SET password = %s WHERE email = %s", (hashed_password, email))
                    connection.commit()

                    # ✅ Success message — JavaScript will handle redirection
                    return jsonify({'message': 'Password reset successful!'}), 200
                else:
                    return jsonify({'error': 'Invalid OTP'}), 400

            else:
                return jsonify({'error': 'Failed to connect to the database'}), 500

        except Exception as e:
            return jsonify({'error': f"Database Error: {e}"}), 500

        finally:
            if connection:
                cursor.close()
                connection.close()




ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
            
#add_question
# Route to render the HTML form and handle form submissions
@app.route('/add_question', methods=['GET', 'POST'])
def add_question():
    if not check_access():  # If hasAccess is not 'true', redirect to login
        return redirect(url_for('login'))

    connection = create_connection()
    cursor = connection.cursor()

    if request.method == 'GET':
        # Render the 'add_question.html' template with existing questions
        return render_template('add_question.html')

    if request.method == 'POST':
        try:

             # Extract data from form fields
            question_text = request.form.get("questionText", "").strip().lower()
            branch = request.form.get("branch", "").strip()
            semester = request.form.get("semester", "").strip()
            subject = request.form.get("subject", "").strip()
            rbt_level = request.form.get("rbtLevel", "").strip()
            co = request.form.get("co", "").strip()
            pi = request.form.get("pi", "").strip()
            marks = request.form.get("marks", "").strip()
            user_id = session['user_id']

             # ✅ Save image if uploaded
            image_path = None
            if 'file' in request.files:
               image_file = request.files['file']
               if image_file and image_file.filename:
                   filename = secure_filename(image_file.filename)
                   image_path = os.path.join(UPLOAD_FOLDER, filename)
                   image_file.save(image_path)

            # Get or insert subject_id
            cursor.execute("SELECT subject_id FROM subjects WHERE subject_name = %s AND branch = %s AND semester = %s",
                           (subject, branch, semester))
            subject_result = cursor.fetchone()

            if not subject_result:
                 return jsonify({"success": False, "message": "Subject doesn't exist!"}), 404  # Return a 404 status
            else:
                subject_id = subject_result[0]

            # Insert new question
            cursor.execute("""
                INSERT INTO questions (question_text, rbt_level, co, pi, marks, subject_id, user_id, image_path) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (question_text, rbt_level, co, pi, marks, subject_id, user_id, image_path))
            connection.commit()

            return jsonify({"success": True, "message": "Question added successfully!"}), 200

        except Exception as e:
            # print(f"Error: {e}")
            connection.rollback()
            return jsonify({"success": False, "message": "An error occurred while adding the question."}), 500

        finally:
            cursor.close()
            connection.close()
            
#For admin View           
@app.route('/get_all_subjects', methods=['GET'])
def get_all_subjects():
    connection = create_connection()
    cursor = connection.cursor()

    # Fetch all subjects
    cursor.execute("SELECT subject_id, subject_name FROM subjects")
    
    subjects = cursor.fetchall()
    connection.close()

    # Format subjects into JSON
    subjects_data = [{'id': s[0], 'name': s[1]} for s in subjects]

    return jsonify({'subjects': subjects_data})


@app.route('/delete_subject/<int:subject_id>', methods=['DELETE'])
def delete_subject(subject_id):
    connection = create_connection()
    cursor = connection.cursor()

    try:
        # Optional: check if subject exists before deleting
        cursor.execute("SELECT * FROM subjects WHERE subject_id = %s", (subject_id,))
        if cursor.fetchone() is None:
            connection.close()
            return jsonify({'error': 'Subject not found'}), 404

        # Delete the subject
        cursor.execute("DELETE FROM subjects WHERE subject_id = %s", (subject_id,))
        connection.commit()
        connection.close()

        return jsonify({'message': f'Subject with ID {subject_id} deleted successfully.'}), 200

    except Exception as e:
        connection.rollback()
        connection.close()
        return jsonify({'error': str(e)}), 500


            
            
@app.route('/fetch_subjects', methods=['GET'])
def fetch_subjects():
    branch = request.args.get('branch')  # Get the branch parameter from the query string
    semester = request.args.get('semester')  # Get the semester parameter from the query string
    # print(branch)
    # print(semester)

    try:
        connection = create_connection()
        cursor = connection.cursor()

        # Query subjects based on branch and semester
        cursor.execute("""
            SELECT subject_name FROM subjects
            WHERE branch = %s AND semester = %s
        """, (branch, semester))

        subjects = cursor.fetchall()  # Fetch subject names
        # print(subjects)

        if subjects:
            # Extract only subject names and return them
            subject_names = [subject[0] for subject in subjects]
            return jsonify(subject_names)  # Return the list of subject names in JSON format
        else:
            return jsonify([])  # Return an empty list if no subjects are found

    except Exception as e:
        return jsonify({'error': f'Error fetching subjects: {str(e)}'}), 500

    finally:
        cursor.close()
        connection.close()




@app.route('/fetch_questions')
def fetch_questions():
    
    if 'user_id' not in session:
        return jsonify([])  # Return empty if not logged in

    user_id = session['user_id']
    connection = create_connection()
    cursor = connection.cursor()

    cursor.execute("SELECT question_text FROM questions WHERE user_id = %s", (user_id,))
    questions = [row['question_text'] for row in cursor.fetchall()]
    
    cursor.close()
    connection.close()
    return jsonify(questions)



@app.route('/get_subjects', methods=['GET'])
def get_subjects():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    user_id = session['user_id']
    # print(user_id)
    connection = create_connection()
    cursor = connection.cursor()

    # Fetch distinct subjects from the questions table based on user_id
    cursor.execute("""
        SELECT DISTINCT s.subject_id, s.subject_name
        FROM subjects s
        JOIN questions q ON s.subject_id = q.subject_id
        WHERE q.user_id = %s
    """, (user_id,))
    
    subjects = cursor.fetchall()
    connection.close()

    # Format subjects into JSON format
    subjects_data = [{'id': s[0], 'name': s[1]} for s in subjects]
    # print(subjects_data)

    return jsonify({'subjects': subjects_data})

@app.route('/show_questions', methods=['GET'])
def show_questions():
    if not check_access():  
        return redirect(url_for('login'))
    
    # Render the template
    return render_template('show_questions.html')

@app.route('/get_questions')
def get_questions():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    user_id = session['user_id']
    subject_id = request.args.get('subject_id')

    if not subject_id:
        return jsonify({'error': 'Subject ID is required'}), 400

    connection = create_connection()
    cursor = connection.cursor()

    # Fetch questions based on user_id and subject_id
    cursor.execute("""
        SELECT question_id, question_text, marks, rbt_level, co, pi, image_path
        FROM questions
        WHERE user_id = %s AND subject_id = %s
    """, (user_id, subject_id))
    
    questions = cursor.fetchall()

    connection.close()

    # Format questions into a list of dictionaries for easy JSON parsing
    questions_data = [
        {
            'id': q[0],
            'text': q[1],
            'marks': q[2],
            'rbt_level': q[3],
            'co': q[4],
            'pi': q[5],
            'image': q[6] if q[6] else None
        }
        for q in questions
    ]

    return jsonify({'questions': questions_data})



@app.route("/update_question", methods=["POST"])
def update_question():
    try:
        # Extract data from form fields
        question_id = request.form.get("question_id")
        question_text = request.form.get("question")
        marks = request.form.get("marks")
        rbt_level = request.form.get("rbt_level")
        co = request.form.get("co")
        pi = request.form.get("pi")

        if not question_id or not question_text:
            return jsonify({"status": "error", "message": "Missing required fields"}), 400

        conn = create_connection()
        cursor = conn.cursor()

        # ✅ Handle Image Upload (if a new image is uploaded)
        if "image" in request.files and request.files["image"].filename:
            image = request.files["image"]
            image_filename = f"question_{question_id}.jpg"
            image_path = os.path.join(app.config["UPLOAD_FOLDER"], image_filename)
            image.save(image_path)

            cursor.execute("""
                UPDATE questions 
                SET question_text = %s, marks = %s, rbt_level = %s, co = %s, pi = %s, image_path = %s 
                WHERE question_id = %s
            """, (question_text, marks, rbt_level, co, pi, image_filename, question_id))
        else:
            cursor.execute("""
                UPDATE questions 
                SET question_text = %s, marks = %s, rbt_level = %s, co = %s, pi = %s 
                WHERE question_id = %s
            """, (question_text, marks, rbt_level, co, pi, question_id))

        conn.commit()
        return jsonify({"status": "success", "message": "Question updated successfully"}), 200

    except Exception as e:
        conn.rollback()  # ✅ Ensure rollback on failure
        return jsonify({"status": "error", "message": f"Failed to update question: {str(e)}"}), 500

    finally:
        conn.close()





@app.route("/delete_question/<int:question_id>", methods=["DELETE"])
def delete_question(question_id):
    conn = create_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("DELETE FROM questions WHERE question_id = %s", (question_id,))
        conn.commit()
        message = "Question deleted successfully."
        status = "success"
    except Exception as e:
        conn.rollback()
        message = f"Failed to delete question: {str(e)}"
        status = "error"
    finally:
        conn.close()

    return jsonify({"message": message, "status": status})



@app.route('/generate_question_paper', methods=['GET', 'POST'])
def generate_question_paper():
    
    if not check_access():  # If 'user_id' is not in session, redirect to login
        flash("Please log in.", "warning")  # Flash a warning message to the user
        return redirect(url_for('login'))  # Redirect to login page

    if request.method == 'POST':
        data = request.get_json()

        branch = data.get('branch')
        semester = data.get('semester')
        subject = data.get('subject')
        subject_code = data.get('subjectCode')
        exam_date = data.get('examDate')
        examination = data.get('examination')
        paper_code = data.get('paperCode')
        total_marks = data.get('totalMarks', 0)

        # print(subject)

        connection = create_connection()
        if not connection:
            return jsonify({
                'status': 'error',
                'message': 'Database connection failed'
            }), 500
        
        try:
            cursor = connection.cursor()

            cursor.execute("SELECT subject_id FROM subjects WHERE subject_name = %s", (subject,))
            subject_result = cursor.fetchone()

            if not subject_result:
                return jsonify({
                    'status': 'error',
                    'message': 'Subject not found in database'
                }), 404

            subject_id = subject_result[0]

            cursor.execute("""
                SELECT question_text, CAST(marks AS SIGNED), rbt_level, co, pi, image_path 
                FROM questions 
                WHERE subject_id = %s
                ORDER BY RAND()
            """, (subject_id,))
            all_questions = cursor.fetchall()

            if not all_questions:
                return jsonify({
                    'status': 'error',
                    'message': 'No questions found for the selected subject'
                }), 404

            cursor.close()

            selected_questions = []
            current_marks = 0

            for question in all_questions:
                if current_marks + question[1] <= total_marks:
                    selected_questions.append(question)
                    current_marks += question[1]
                if current_marks == total_marks:
                    break

            if current_marks != total_marks:
                flash("Could not select exact marks, closest possible paper generated.", 'warning')

            # Create the Word document
            doc = Document()

            header = doc.add_paragraph()
            header.alignment = 1
            header.add_run("D Y Patil College of Engineering, ").bold = True
            header.add_run("Salokhenagar, Kolhapur 416007\n").bold = True
            header.add_run("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING (DATA SCIENCE)\n").bold = True

            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'

            cells = table.rows[0].cells
            cells[0].text = examination
            cells[1].text = f"Sem : {semester}"
            cells[2].text = f"Subject : {subject}"
            cells[3].text = f"Subject Code: {subject_code}"

            cells = table.rows[1].cells
            cells[0].text = f"Date : {exam_date}"
            cells[1].text = f"Maximum Marks: {total_marks}"
            cells[2].text = f"QP Code: {paper_code}"
            cells[3].text = "All questions are compulsory."

            line_paragraph = doc.add_paragraph()
            line_run = line_paragraph.add_run("_" * 95)
            line_run.bold = True
            line_paragraph.alignment = 1

            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            headers = ["Q. No.", "Question", "Marks", "RBT Level", "CO", "PI", "Marks Secured"]
            for i, header_text in enumerate(headers):
                table.rows[0].cells[i].text = header_text

            for i, (question_text, marks, rbt_level, co, pi, image_path) in enumerate(selected_questions, start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = question_text
                row_cells[2].text = str(marks)
                row_cells[3].text = str(rbt_level)
                row_cells[4].text = str(co)
                row_cells[5].text = str(pi)
                row_cells[6].text = ""  # Marks Secured

                if image_path:
                    image_path = image_path.replace('\\', '/')

                    # If the image_path doesn't already have 'static/uploads/', prepend it
                    if not image_path.startswith('static/uploads/'):
                        image_path = 'static/uploads/' + image_path

                    # Check if the image file exists
                    if os.path.exists(image_path):
                        # Insert image into the document in the same table cell (below the question text)
                        try:
                            # Add image to the same cell under the question
                            image_cell = row_cells[1]
                            # Create a new paragraph inside the cell
                            image_paragraph = image_cell.add_paragraph()
                            image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            # Add image below the question text (adjust width as needed)
                            image_paragraph.add_run().add_picture(image_path, width=Inches(1.0))  # You can adjust width here
                        except Exception as e:
                            print(f"Error adding image {image_path}: {e}")
                    else:
                        print(f"Image not found: {image_path}")

            # Save the document to memory and send it as a response
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return send_file(file_stream, as_attachment=True, download_name=f"{subject}.docx")

        except psycopg2.Error as e:
            return jsonify({
                'status': 'error',
                'message': f"Database error: {e}"
            }), 500

        finally:
            if connection:
                connection.close()


    return render_template('generate_question_paper.html')



@app.route('/logout')
def logout():
    # Clear the session to log out the user
    session.clear()
    flash('You have been logged out successfully.', 'info')
    return redirect(url_for('login'))


create_database() 
# Initialize database and tables
initialize_database()

if __name__ == "__main__":
    app.run(debug=True)
    

